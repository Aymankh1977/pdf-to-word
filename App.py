import io
import re
import os
import zipfile
import streamlit as st
import pdfplumber
import pikepdf
from pypdf import PdfReader, PdfWriter
from pdf2image import convert_from_bytes
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import Color, HexColor

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PDF Studio",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,300;0,400;1,300;1,400&family=Jost:wght@300;400;500;600&display=swap');

html, body, [class*="css"] { font-family: 'Jost', sans-serif !important; }
.stApp { background-color: #f7f3ee !important; }

/* ── Hero ── */
.hero { text-align:center; padding:2.5rem 1rem 1.2rem; margin-bottom:0; }
.hero-eyebrow { font-size:0.65rem; font-weight:500; letter-spacing:0.28em; text-transform:uppercase; color:#a89880; margin-bottom:0.6rem; }
.hero h1 { font-family:'Cormorant Garamond',serif !important; font-size:3rem !important; font-weight:300 !important; color:#1c1917 !important; letter-spacing:-0.02em !important; margin-bottom:0.4rem !important; }
.hero h1 em { font-style:italic; color:#8b5e52; }
.hero-sub { font-size:0.76rem; color:#9a8e82; font-weight:300; letter-spacing:0.06em; }

/* ── Navigation bar ── */
.nav-wrap { background:#fff; border-bottom:2px solid #e2d9ce; padding:16px 8px 12px; margin-bottom:1.5rem; box-shadow:0 2px 8px rgba(28,25,23,0.05); }
.nav-group-label {
    font-size:0.6rem !important; font-weight:600 !important;
    letter-spacing:0.22em !important; text-transform:uppercase !important;
    color:#a89880 !important; margin:0 0 6px 0 !important;
    padding:0 !important; line-height:1 !important;
}
/* Active nav button — override stButton for nav buttons only */
[data-testid="stHorizontalBlock"] .stButton > button,
.nav-wrap .stButton > button {
    background: #f5ede8 !important;
    color: #6b5f55 !important;
    border: 1px solid #e2d9ce !important;
    border-radius: 4px !important;
    font-size: 0.72rem !important;
    font-weight: 400 !important;
    letter-spacing: 0.04em !important;
    text-transform: none !important;
    padding: 0.4rem 0.6rem !important;
    margin-bottom: 4px !important;
    text-align: left !important;
    transition: all 0.15s !important;
}
[data-testid="stHorizontalBlock"] .stButton > button:hover,
.nav-wrap .stButton > button:hover {
    background: #ede0d8 !important;
    color: #1c1917 !important;
    border-color: #c4b0a6 !important;
}

/* ── Tool heading ── */
.tool-heading { display:flex; align-items:baseline; gap:0.8rem; flex-wrap:wrap; margin-bottom:1.8rem; padding-bottom:1rem; border-bottom:1px solid #e2d9ce; }
.tool-heading h2 { font-family:'Cormorant Garamond',serif !important; font-size:1.9rem !important; font-weight:400 !important; color:#1c1917 !important; margin:0 !important; }
.tool-heading .tool-tag { font-size:0.65rem; font-weight:500; letter-spacing:0.14em; text-transform:uppercase; color:#a89880; background:#ede6dc; padding:3px 10px; border-radius:20px; }

/* ── File uploader ── */
[data-testid="stFileUploader"] { background:#fff !important; border:1px solid #d5ccc4 !important; border-radius:3px !important; }

/* ── Buttons ── */
.stButton > button {
    background:#1c1917 !important; color:#f7f3ee !important; border:none !important;
    border-radius:2px !important; font-family:'Jost',sans-serif !important;
    font-weight:500 !important; font-size:0.72rem !important;
    letter-spacing:0.14em !important; text-transform:uppercase !important;
    padding:0.65rem 2rem !important; transition:background 0.2s !important;
}
.stButton > button:hover { background:#3d3530 !important; }
[data-testid="stDownloadButton"] > button {
    background:transparent !important; color:#1c1917 !important;
    border:1.5px solid #1c1917 !important; border-radius:2px !important;
    font-family:'Jost',sans-serif !important; font-weight:500 !important;
    font-size:0.72rem !important; letter-spacing:0.14em !important;
    text-transform:uppercase !important; width:100% !important; transition:all 0.2s !important;
}
[data-testid="stDownloadButton"] > button:hover { background:#1c1917 !important; color:#f7f3ee !important; }

/* ── Alerts ── */
[data-testid="stInfo"]    { background:#ede6dc !important; border:none !important; border-left:3px solid #a89880 !important; border-radius:0 !important; }
[data-testid="stSuccess"] { background:#e8ede6 !important; border:none !important; border-left:3px solid #6b8b5e !important; border-radius:0 !important; }
[data-testid="stError"]   { background:#f0e6e6 !important; border:none !important; border-left:3px solid #8b5e5e !important; border-radius:0 !important; }

/* ── Result card ── */
.result-card { background:#fff; border:1px solid #e2d9ce; border-left:3px solid #8b5e52; padding:1rem 1.4rem; margin-bottom:0.6rem; border-radius:0 2px 2px 0; }
.result-card .fname { font-weight:500; font-size:0.88rem; color:#1c1917; margin-bottom:3px; }
.result-card .fmeta { font-size:0.76rem; color:#9a8e82; }

/* ── Inputs ── */
.stTextInput input, .stNumberInput input, .stTextArea textarea {
    background:#fff !important; border:1px solid #d5ccc4 !important;
    border-radius:2px !important; color:#1c1917 !important; font-family:'Jost',sans-serif !important;
}
.stTextInput input:focus, .stNumberInput input:focus, .stTextArea textarea:focus { border-color:#8b5e52 !important; }
[data-baseweb="select"] > div { background:#fff !important; border:1px solid #d5ccc4 !important; border-radius:2px !important; }
label[data-testid="stWidgetLabel"] p { font-size:0.75rem !important; font-weight:500 !important; letter-spacing:0.08em !important; text-transform:uppercase !important; color:#6b5f55 !important; }

/* ── Progress ── */
.stProgress > div > div { background:#8b5e52 !important; }
hr { border-color:#e2d9ce !important; }
.stCaption p { font-size:0.73rem !important; color:#a89880 !important; font-style:italic !important; }
#MainMenu, footer, [data-testid="stToolbar"] { visibility:hidden !important; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════
#  UTILITY FUNCTIONS
# ═══════════════════════════════════════════════════════════════════

def get_pdf_info(pdf_bytes):
    r = PdfReader(io.BytesIO(pdf_bytes))
    encrypted = r.is_encrypted
    if encrypted:
        # Can't count pages without the password — return safely
        return {"pages": 0, "encrypted": True}
    return {"pages": len(r.pages), "encrypted": False}

def pdf_page_previews(pdf_bytes, dpi=80, max_pages=20):
    """Return list of PIL images for preview."""
    try:
        imgs = convert_from_bytes(pdf_bytes, dpi=dpi, first_page=1, last_page=max_pages)
        return imgs
    except Exception:
        return []

# ── 1. PDF → Word ─────────────────────────────────────────────────

def clean_text(text):
    if not text: return ""
    text = text.replace('\xa0', ' ')
    return re.sub(r'[ \t]{2,}', ' ', text).strip()

def looks_like_heading(text, font_size=None, is_bold=False):
    s = text.strip()
    if not s: return 0
    if s.isupper() and 3 <= len(s) <= 80: return 1
    if is_bold and len(s) <= 80 and not s.endswith('.'): return 2
    if font_size:
        if font_size >= 18: return 1
        if font_size >= 14: return 2
        if font_size >= 12 and is_bold: return 3
    return 0

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),'AAAAAA')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_table_to_doc(doc, table_data, font_size_pt=11):
    rows = [r for r in table_data if any(c for c in r)]
    if not rows: return
    num_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(num_cols):
            ct = row[ci] if ci < len(row) else ""
            cell = tbl.cell(ri, ci)
            cell.text = clean_text(ct or "")
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(font_size_pt - 1)
                if ri == 0: run.bold = True
            if ri == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'E8E8E8')
                tcPr.append(shd)
    set_table_border(tbl)
    doc.add_paragraph()

def extract_page_content(page):
    content = []
    tables = page.extract_tables()
    try: table_bboxes = [t.bbox for t in page.find_tables()]
    except: table_bboxes = []
    try: image_bboxes = [(i['x0'],i['top'],i['x1'],i['bottom']) for i in page.images]
    except: image_bboxes = []
    words = page.extract_words(extra_attrs=["size","fontname"])
    if not words:
        raw = page.extract_text()
        if raw:
            for line in raw.splitlines():
                line = clean_text(line)
                if line: content.append({"type":"paragraph","text":line,"font_size":None,"bold":False})
        for tbl in tables:
            if tbl: content.append({"type":"table","data":tbl})
        return content, image_bboxes
    lines = {}
    for w in words: lines.setdefault(round(w['top'],1),[]).append(w)
    def in_r(y,bbs): return any(b[1]<=y<=b[3] for b in bbs)
    prev_bottom = None
    for top, lw in sorted(lines.items()):
        if in_r(top,table_bboxes) or in_r(top,image_bboxes): continue
        lw.sort(key=lambda w: w['x0'])
        lt = clean_text(' '.join(w['text'] for w in lw))
        if not lt: continue
        sizes = [w.get('size',0) for w in lw if w.get('size')]
        fs = max(sizes) if sizes else None
        bold = any('Bold' in w.get('fontname','') for w in lw)
        if prev_bottom and (top-prev_bottom)>12: content.append({"type":"blank"})
        content.append({"type":"paragraph","text":lt,"font_size":fs,"bold":bold})
        prev_bottom = top + lw[0].get('height',10)
    for tbl in tables:
        if tbl: content.append({"type":"table","data":tbl})
    return content, image_bboxes

def crop_img(page_img, bbox, pw, ph):
    iw, ih = page_img.size
    sx, sy = iw/pw, ih/ph
    x0,y0,x1,y1 = int(bbox[0]*sx),int(bbox[1]*sy),int(bbox[2]*sx),int(bbox[3]*sy)
    pad=4; x0,y0=max(0,x0-pad),max(0,y0-pad); x1,y1=min(iw,x1+pad),min(ih,y1+pad)
    if x1<=x0 or y1<=y0: return None
    return page_img.crop((x0,y0,x1,y1))

def convert_pdf_to_docx(pdf_bytes, font_size=11, include_images=True, ocr_fallback=True, dpi=150):
    doc = Document()
    for s in doc.sections:
        s.top_margin=s.bottom_margin=Inches(1)
        s.left_margin=s.right_margin=Inches(1.2)
    doc.styles['Normal'].font.name='Calibri'
    doc.styles['Normal'].font.size=Pt(font_size)
    page_images=[]
    if include_images:
        try: page_images=convert_from_bytes(pdf_bytes,dpi=dpi)
        except: pass
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        total=len(pdf.pages)
        for pi, page in enumerate(pdf.pages):
            content, image_bboxes = extract_page_content(page)
            raw_text = page.extract_text() or ""
            if not raw_text.strip() and ocr_fallback and page_images:
                try:
                    import pytesseract
                    if pi < len(page_images):
                        for line in pytesseract.image_to_string(page_images[pi]).splitlines():
                            line=clean_text(line)
                            if line: content.append({"type":"paragraph","text":line,"font_size":None,"bold":False})
                except: pass
            for item in content:
                if item["type"]=="blank": continue
                elif item["type"]=="paragraph":
                    text=item["text"]
                    level=looks_like_heading(text,item.get("font_size"),item.get("bold",False))
                    if level in (1,2,3):
                        h=doc.add_heading(text,level=level)
                        h.runs[0].font.size=Pt(font_size+(6-level*2))
                    else:
                        p=doc.add_paragraph(text)
                        for run in p.runs:
                            run.font.size=Pt(font_size)
                            if item.get("bold"): run.bold=True
                elif item["type"]=="table":
                    add_table_to_doc(doc,item["data"],font_size)
            if include_images and page_images and pi<len(page_images):
                pg_img=page_images[pi]; used=[]
                for bbox in image_bboxes:
                    w,h=bbox[2]-bbox[0],bbox[3]-bbox[1]
                    if w<20 or h<20: continue
                    if any(not(bbox[2]<u[0] or bbox[0]>u[2] or bbox[3]<u[1] or bbox[1]>u[3]) for u in used): continue
                    used.append(bbox)
                    cropped=crop_img(pg_img,bbox,page.width,page.height)
                    if not cropped: continue
                    ib=io.BytesIO(); cropped.save(ib,format='PNG'); ib.seek(0)
                    iw2,ih2=cropped.size; mw=Inches(5)
                    asp=ih2/iw2 if iw2>0 else 1; dw=min(mw,Inches(iw2/dpi))
                    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(ib,width=dw); doc.add_paragraph()
            if pi<total-1: doc.add_page_break()
    out=io.BytesIO(); doc.save(out); return out.getvalue()

# ── 2. Merge PDFs ─────────────────────────────────────────────────

def merge_pdfs(pdf_list):
    writer = PdfWriter()
    for pdf_bytes in pdf_list:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()

# ── 3. Split PDF ──────────────────────────────────────────────────

def split_pdf(pdf_bytes, ranges):
    """ranges: list of (start, end) 1-indexed inclusive tuples"""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    results = []
    for start, end in ranges:
        writer = PdfWriter()
        for i in range(start-1, min(end, len(reader.pages))):
            writer.add_page(reader.pages[i])
        out = io.BytesIO()
        writer.write(out)
        results.append((f"pages_{start}_to_{end}.pdf", out.getvalue()))
    return results

# ── 4. Rotate pages ───────────────────────────────────────────────

def rotate_pdf(pdf_bytes, angle, page_nums=None):
    """angle: 90, 180, 270. page_nums: list of 1-indexed, or None for all."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for i, page in enumerate(reader.pages):
        if page_nums is None or (i+1) in page_nums:
            page.rotate(angle)
        writer.add_page(page)
    out = io.BytesIO(); writer.write(out); return out.getvalue()

# ── 5. Watermark ──────────────────────────────────────────────────

def add_watermark(pdf_bytes, text, opacity=0.25, font_size=60, color_hex="#cc3333", angle=45):
    r = int(color_hex[1:3],16)/255
    g = int(color_hex[3:5],16)/255
    b = int(color_hex[5:7],16)/255

    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        # Build a watermark sized to THIS page
        pw = float(page.mediabox.width)
        ph = float(page.mediabox.height)
        wm_buf = io.BytesIO()
        c = rl_canvas.Canvas(wm_buf, pagesize=(pw, ph))
        c.setFillColor(Color(r, g, b, alpha=opacity))
        c.setFont("Helvetica-Bold", font_size)
        c.saveState()
        c.translate(pw/2, ph/2)
        c.rotate(angle)
        c.drawCentredString(0, 0, text)
        c.restoreState()
        c.save()
        wm_page = PdfReader(io.BytesIO(wm_buf.getvalue())).pages[0]
        page.merge_page(wm_page)
        writer.add_page(page)
    out = io.BytesIO(); writer.write(out); return out.getvalue()

# ── 6. Password protect / unlock ─────────────────────────────────

def protect_pdf(pdf_bytes, user_password, owner_password=None):
    writer = PdfWriter()
    reader = PdfReader(io.BytesIO(pdf_bytes))
    for page in reader.pages: writer.add_page(page)
    writer.encrypt(user_password, owner_password or user_password)
    out = io.BytesIO(); writer.write(out); return out.getvalue()

def unlock_pdf(pdf_bytes, password):
    reader = PdfReader(io.BytesIO(pdf_bytes))
    if reader.is_encrypted:
        result = reader.decrypt(password)
        # decrypt returns 0 if password is wrong
        if result == 0:
            raise ValueError("Incorrect password")
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()

# ── 7. Compress ───────────────────────────────────────────────────

def compress_pdf(pdf_bytes, quality="medium"):
    """Compress PDF. quality: 'high' (light), 'medium', 'low' (aggressive)."""
    # Try Ghostscript-style compression via pikepdf image recompression
    out = io.BytesIO()
    settings = {
        "high":   {"recompress": True},
        "medium": {"recompress": True},
        "low":    {"recompress": True},
    }
    try:
        with pikepdf.open(io.BytesIO(pdf_bytes)) as pk:
            # Recompress images inside the PDF
            for page in pk.pages:
                for name, raw in list(page.images.items()):
                    try:
                        pdfimg = pikepdf.PdfImage(raw)
                        pil = pdfimg.as_pil_image()
                        # Downsample based on quality
                        if quality == "low":
                            max_dim = 800
                            jpg_q = 40
                        elif quality == "medium":
                            max_dim = 1200
                            jpg_q = 60
                        else:
                            max_dim = 2000
                            jpg_q = 80
                        w, h = pil.size
                        if max(w, h) > max_dim:
                            ratio = max_dim / max(w, h)
                            pil = pil.resize((int(w*ratio), int(h*ratio)))
                        # Re-encode as JPEG
                        img_buf = io.BytesIO()
                        pil.convert("RGB").save(img_buf, "JPEG", quality=jpg_q)
                        raw.write(img_buf.getvalue(), filter=pikepdf.Name("/DCTDecode"))
                    except Exception:
                        continue
            pk.save(out, compress_streams=True, recompress_flate=True)
        return out.getvalue()
    except Exception:
        # Fallback to basic compression
        out2 = io.BytesIO()
        with pikepdf.open(io.BytesIO(pdf_bytes)) as pk:
            pk.save(out2, compress_streams=True, recompress_flate=True)
        return out2.getvalue()

# ── 8. Extract pages as images ────────────────────────────────────

def extract_as_images(pdf_bytes, dpi=150, fmt="PNG"):
    imgs = convert_from_bytes(pdf_bytes, dpi=dpi)
    results = []
    for i, img in enumerate(imgs):
        buf = io.BytesIO()
        img.save(buf, format=fmt)
        results.append((f"page_{i+1}.{fmt.lower()}", buf.getvalue()))
    return results

# ── 9. Add text annotation ────────────────────────────────────────

def add_text_annotation(pdf_bytes, text, page_num=1, x=100, y=100,
                         font_size=14, color_hex="#cc0000"):
    annot_buf = io.BytesIO()
    c = rl_canvas.Canvas(annot_buf, pagesize=letter)
    r2 = int(color_hex[1:3],16)/255
    g2 = int(color_hex[3:5],16)/255
    b2 = int(color_hex[5:7],16)/255
    c.setFillColorRGB(r2, g2, b2)
    c.setFont("Helvetica", font_size)
    c.drawString(x, y, text)
    c.save()
    annot_page = PdfReader(io.BytesIO(annot_buf.getvalue())).pages[0]

    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for i, page in enumerate(reader.pages):
        if i+1 == page_num:
            page.merge_page(annot_page)
        writer.add_page(page)
    out = io.BytesIO(); writer.write(out); return out.getvalue()

# ── 10. Redact ────────────────────────────────────────────────────

def redact_pdf(pdf_bytes, page_num=1, x=100, y=100, width=200, height=20):
    """TRUE redaction: rasterize the page so content underneath is destroyed,
    then draw the black box. Text under the box cannot be extracted."""
    import subprocess, tempfile, glob
    from PIL import Image as PILImage, ImageDraw
    from reportlab.lib.utils import ImageReader

    reader = PdfReader(io.BytesIO(pdf_bytes))
    pw = float(reader.pages[page_num-1].mediabox.width)
    ph = float(reader.pages[page_num-1].mediabox.height)

    # Render the target page to an image at 150 DPI
    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = os.path.join(tmp, "in.pdf")
        with open(pdf_path, "wb") as fo:
            fo.write(pdf_bytes)
        prefix = os.path.join(tmp, "pg")
        subprocess.run(["pdftoppm", "-r", "150", "-png",
                        "-f", str(page_num), "-l", str(page_num), pdf_path, prefix],
                       capture_output=True, timeout=60)
        matches = glob.glob(prefix + "*")
        if not matches:
            raise RuntimeError("Could not render page for redaction")
        page_img = PILImage.open(matches[0]).convert("RGB")

    # Draw black box on the rasterized image (destroys content underneath)
    img_w, img_h = page_img.size
    scale_x = img_w / pw
    scale_y = img_h / ph
    # PDF y is from bottom; image y is from top
    box_left   = x * scale_x
    box_top    = (ph - y - height) * scale_y
    box_right  = (x + width) * scale_x
    box_bottom = (ph - y) * scale_y
    draw = ImageDraw.Draw(page_img)
    draw.rectangle([box_left, box_top, box_right, box_bottom], fill=(0, 0, 0))

    # Rebuild that page from the redacted image
    img_buf = io.BytesIO()
    page_img.save(img_buf, "PNG")
    img_buf.seek(0)
    new_page_buf = io.BytesIO()
    c = rl_canvas.Canvas(new_page_buf, pagesize=(pw, ph))
    c.drawImage(ImageReader(img_buf), 0, 0, width=pw, height=ph)
    c.save()
    redacted_page = PdfReader(io.BytesIO(new_page_buf.getvalue())).pages[0]

    writer = PdfWriter()
    for i, page in enumerate(reader.pages):
        if i+1 == page_num:
            writer.add_page(redacted_page)
        else:
            writer.add_page(page)
    out = io.BytesIO(); writer.write(out); return out.getvalue()

# ── 11. Reorder pages ─────────────────────────────────────────────

def reorder_pages(pdf_bytes, new_order):
    """new_order: list of 1-indexed page numbers in desired order."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for n in new_order:
        if 1 <= n <= len(reader.pages):
            writer.add_page(reader.pages[n-1])
    out = io.BytesIO(); writer.write(out); return out.getvalue()

# ═══════════════════════════════════════════════════════════════════
#  UI
# ═══════════════════════════════════════════════════════════════════

TOOLS = [
    ("📝", "PDF → Word",       "PDF to Word document"),
    ("📊", "PDF → Excel",      "PDF to Excel spreadsheet"),
    ("📋", "PDF → CSV",        "PDF to CSV data file"),
    ("🖼️", "PDF → Images",     "PDF pages to PNG/JPG"),
    ("📄", "PDF → Text",       "PDF to plain text"),
    ("📝", "PDF → Markdown",   "PDF to Markdown"),
    ("📝", "Word → Markdown",  "DOCX to Markdown"),
    ("📝", "Excel → Markdown", "Excel to Markdown tables"),
    ("🔗", "Merge PDFs",       "Combine multiple PDFs"),
    ("✂️",  "Split PDF",        "Extract page ranges"),
    ("🔄", "Rotate Pages",     "Rotate any page"),
    ("💧", "Watermark",        "Stamp text on pages"),
    ("🗜️", "Compress",         "Reduce file size"),
    ("🔒", "Protect / Unlock", "Password management"),
    ("✍️",  "Add Text",         "Annotate pages"),
    ("⬛", "Redact",           "Black-out content"),
    ("📋", "Reorder Pages",    "Drag pages into order"),
    ("🖼️", "Images → PDF",     "Images to PDF"),
    ("📄", "Word → PDF",       "Word docs to PDF"),
    ("📊", "Excel → PDF",      "Spreadsheets to PDF"),
]

st.markdown("""
<div class="hero">
    <div class="hero-eyebrow">Professional Document Tools</div>
    <h1>PDF <em>Studio</em></h1>
    <div class="hero-sub">Convert &nbsp;&middot;&nbsp; Merge &nbsp;&middot;&nbsp; Split &nbsp;&middot;&nbsp; Watermark &nbsp;&middot;&nbsp; Compress &nbsp;&middot;&nbsp; Protect &nbsp;&middot;&nbsp; Annotate</div>
</div>
""", unsafe_allow_html=True)

# ── Tool navigation ──────────────────────────────────────────────
TOOL_GROUPS = {
    "📤  PDF to...":  ["PDF → Word", "PDF → Excel", "PDF → CSV",
                       "PDF → Images", "PDF → Text",
                       "PDF → Markdown", "Word → Markdown", "Excel → Markdown"],
    "📥  ...to PDF":  ["Images → PDF", "Word → PDF", "Excel → PDF"],
    "✏️  Edit PDF":   ["Add Text", "Redact", "Watermark", "Rotate Pages"],
    "📂  Manage PDF": ["Merge PDFs", "Split PDF", "Reorder Pages",
                       "Compress", "Protect / Unlock"],
}

# Build flat ordered list matching TOOLS order for display
ALL_TOOL_NAMES = [t[1] for t in TOOLS]

# Nav bar using columns + buttons
st.markdown('<div class="nav-wrap">', unsafe_allow_html=True)
nav_cols = st.columns(len(TOOL_GROUPS))
for col_idx, (group_label, items) in enumerate(TOOL_GROUPS.items()):
    with nav_cols[col_idx]:
        st.markdown(f'<p class="nav-group-label">{group_label}</p>', unsafe_allow_html=True)
        for item in items:
            icon = next((t[0] for t in TOOLS if t[1] == item), "")
            is_active = st.session_state.get("selected_tool", TOOLS[0][1]) == item
            btn_style = "nav-btn-active" if is_active else "nav-btn"
            if st.button(f"{icon}  {item}", key=f"nav_{item}",
                         use_container_width=True):
                st.session_state.selected_tool = item
                st.rerun()
st.markdown('</div>', unsafe_allow_html=True)
st.markdown("---")

if "selected_tool" not in st.session_state:
    st.session_state.selected_tool = TOOLS[0][1]
selected_tool = st.session_state.selected_tool


icon = [t[0] for t in TOOLS if t[1]==selected_tool][0]
desc = [t[2] for t in TOOLS if t[1]==selected_tool][0]
st.markdown(f'''<div class="tool-heading">
    <span class="tool-icon">{icon}</span>
    <h2>{selected_tool}</h2>
    <span class="tool-tag">{desc}</span>
</div>''', unsafe_allow_html=True)

# ── PDF → Word ────────────────────────────────────────────────────
if selected_tool == "PDF → Word":
    col1, col2 = st.columns([2,1])
    with col1:
        files = st.file_uploader("Upload PDF(s)", type=["pdf"], accept_multiple_files=True)
    with col2:
        font_size = st.slider("Font size", 9, 14, 11)
        include_images = st.checkbox("Extract figures", value=True)
        ocr = st.checkbox("OCR for scanned PDFs", value=True)
        dpi = st.select_slider("Image DPI", [72,100,150,200], value=150)
        as_zip = st.checkbox("Download as ZIP", value=False)

    if files and st.button("Convert to Word"):
        results = []
        bar = st.progress(0)
        for idx, f in enumerate(files):
            with st.spinner(f"Converting {f.name}…"):
                try:
                    docx_bytes = convert_pdf_to_docx(f.read(), font_size, include_images, ocr, dpi)
                    out_name = os.path.splitext(f.name)[0] + ".docx"
                    results.append({"name":f.name,"out":out_name,"bytes":docx_bytes,"ok":True})
                except Exception as e:
                    results.append({"name":f.name,"error":str(e),"ok":False})
            bar.progress((idx+1)/len(files))
        bar.empty()

        if as_zip and any(r["ok"] for r in results):
            zb = io.BytesIO()
            with zipfile.ZipFile(zb,'w',zipfile.ZIP_DEFLATED) as zf:
                for r in results:
                    if r["ok"]: zf.writestr(r["out"], r["bytes"])
            st.download_button("⬇ Download ZIP", zb.getvalue(), "converted.zip", "application/zip")
        else:
            for r in results:
                status = "✅ Converted" if r["ok"] else f"❌ {r.get('error','')}"
                st.markdown(f'<div class="result-card"><div class="fname">📄 {r["name"]}</div><div class="fmeta">{status}</div></div>', unsafe_allow_html=True)
                if r["ok"]:
                    st.download_button(f"⬇ Download {r['out']}", r["bytes"], r["out"],
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=r["out"])

# ── Merge ─────────────────────────────────────────────────────────
elif selected_tool == "Merge PDFs":
    files = st.file_uploader("Upload PDFs to merge (in order)", type=["pdf"], accept_multiple_files=True)
    out_name = st.text_input("Output filename", value="merged.pdf")
    if files:
        st.info(f"{len(files)} files selected — they will be merged in the order shown above.")
        if st.button("Merge PDFs"):
            with st.spinner("Merging…"):
                try:
                    result = merge_pdfs([f.read() for f in files])
                    info = get_pdf_info(result)
                    st.success(f"✅ Merged {len(files)} files → {info['pages']} pages")
                    st.download_button("⬇ Download Merged PDF", result, out_name, "application/pdf")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Split ─────────────────────────────────────────────────────────
elif selected_tool == "Split PDF":
    f = st.file_uploader("Upload PDF", type=["pdf"])
    if f:
        pdf_bytes = f.read()
        info = get_pdf_info(pdf_bytes)
        st.info(f"📄 {info['pages']} pages total")

        split_mode = st.radio("Split mode", ["Every page (individual files)", "Custom ranges"])
        if split_mode == "Every page (individual files)":
            ranges = [(i,i) for i in range(1, info['pages']+1)]
        else:
            range_input = st.text_input("Page ranges (e.g. 1-3, 4-6, 7-10)", value=f"1-{info['pages']}")
            ranges = []
            for part in range_input.split(','):
                part = part.strip()
                if '-' in part:
                    a,b = part.split('-')
                    ranges.append((int(a.strip()), int(b.strip())))
                elif part.isdigit():
                    ranges.append((int(part), int(part)))

        if st.button("Split PDF"):
            with st.spinner("Splitting…"):
                try:
                    results = split_pdf(pdf_bytes, ranges)
                    st.success(f"✅ Created {len(results)} file(s)")
                    if len(results) == 1:
                        st.download_button(f"⬇ Download {results[0][0]}", results[0][1], results[0][0], "application/pdf")
                    else:
                        zb = io.BytesIO()
                        with zipfile.ZipFile(zb,'w') as zf:
                            for name, data in results: zf.writestr(name, data)
                        st.download_button("⬇ Download All as ZIP", zb.getvalue(), "split_pages.zip", "application/zip")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Rotate ────────────────────────────────────────────────────────
elif selected_tool == "Rotate Pages":
    f = st.file_uploader("Upload PDF", type=["pdf"])
    if f:
        pdf_bytes = f.read()
        info = get_pdf_info(pdf_bytes)
        st.info(f"📄 {info['pages']} pages")
        col1, col2 = st.columns(2)
        with col1:
            angle = st.selectbox("Rotation angle", [90, 180, 270], format_func=lambda x: f"{x}°")
        with col2:
            page_mode = st.radio("Apply to", ["All pages", "Specific pages"])
        page_nums = None
        if page_mode == "Specific pages":
            pg_input = st.text_input("Page numbers (e.g. 1, 3, 5)", value="1")
            page_nums = [int(p.strip()) for p in pg_input.split(',') if p.strip().isdigit()]

        if st.button("Rotate"):
            with st.spinner("Rotating…"):
                try:
                    result = rotate_pdf(pdf_bytes, angle, page_nums)
                    st.success("✅ Done!")
                    # Preview first rotated page
                    try:
                        import subprocess, tempfile, glob
                        from PIL import Image as PILImg
                        with tempfile.TemporaryDirectory() as tmp:
                            pp = os.path.join(tmp, "r.pdf")
                            with open(pp,"wb") as fo: fo.write(result)
                            pref = os.path.join(tmp, "pg")
                            first_pg = page_nums[0] if page_nums else 1
                            subprocess.run(["pdftoppm","-r","80","-png","-f",str(first_pg),
                                            "-l",str(first_pg),pp,pref], capture_output=True, timeout=30)
                            m = glob.glob(pref+"*")
                            if m:
                                st.image(PILImg.open(m[0]), caption=f"Rotated page {first_pg}", width=300)
                    except Exception:
                        pass
                    out_name = os.path.splitext(f.name)[0] + f"_rotated.pdf"
                    st.download_button("⬇ Download", result, out_name, "application/pdf")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Watermark ─────────────────────────────────────────────────────
elif selected_tool == "Watermark":
    f = st.file_uploader("Upload PDF", type=["pdf"])
    if f:
        col1, col2 = st.columns(2)
        with col1:
            wm_text = st.text_input("Watermark text", value="CONFIDENTIAL")
            font_size_wm = st.slider("Font size", 20, 100, 60)
            angle = st.slider("Angle (degrees)", 0, 90, 45)
        with col2:
            opacity = st.slider("Opacity", 0.05, 0.8, 0.25)
            color = st.color_picker("Colour", "#cc3333")

        if st.button("Add Watermark"):
            with st.spinner("Adding watermark…"):
                try:
                    result = add_watermark(f.read(), wm_text, opacity, font_size_wm, color, angle)
                    st.success("✅ Watermark added!")
                    out_name = os.path.splitext(f.name)[0] + "_watermarked.pdf"
                    st.download_button("⬇ Download", result, out_name, "application/pdf")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Compress ──────────────────────────────────────────────────────
elif selected_tool == "Compress":
    files = st.file_uploader("Upload PDF(s)", type=["pdf"], accept_multiple_files=True)
    quality = st.select_slider(
        "Compression level",
        options=["high", "medium", "low"],
        value="medium",
        format_func=lambda x: {"high":"High quality (light compression)",
                                "medium":"Balanced",
                                "low":"Smallest size (aggressive)"}[x]
    )
    if files and st.button("Compress"):
        bar = st.progress(0)
        for idx, f in enumerate(files):
            with st.spinner(f"Compressing {f.name}…"):
                try:
                    original = f.read()
                    compressed = compress_pdf(original, quality)
                    saving = (1 - len(compressed)/len(original)) * 100
                    saving_txt = f"{saving:.1f}% smaller" if saving > 0 else "already optimised"
                    st.markdown(f'<div class="result-card"><div class="fname">🗜️ {f.name}</div><div class="fmeta">{len(original):,} bytes → {len(compressed):,} bytes ({saving_txt})</div></div>', unsafe_allow_html=True)
                    out_name = os.path.splitext(f.name)[0] + "_compressed.pdf"
                    st.download_button(f"⬇ Download {out_name}", compressed, out_name, "application/pdf", key=out_name)
                except Exception as e:
                    st.error(f"❌ {f.name}: {e}")
            bar.progress((idx+1)/len(files))
        bar.empty()

# ── Protect / Unlock ──────────────────────────────────────────────
elif selected_tool == "Protect / Unlock":
    f = st.file_uploader("Upload PDF", type=["pdf"])
    if f:
        pdf_bytes = f.read()
        info = get_pdf_info(pdf_bytes)

        if info["encrypted"]:
            st.info("🔒 This PDF is password-protected. Enter the password to unlock it, or add a new password.")
        else:
            st.info("🔓 This PDF is not protected. You can add a password to secure it.")

        # Offer both actions regardless, but default based on state
        default_idx = 1 if info["encrypted"] else 0
        action = st.radio("Action", ["🔒 Add Password", "🔓 Remove Password"],
                          index=default_idx, horizontal=True)

        if "Add Password" in action:
            if info["encrypted"]:
                st.warning("This PDF is already encrypted. Remove the existing password first before adding a new one.")
            col1, col2 = st.columns(2)
            with col1:
                user_pw = st.text_input("Password (required to open)", type="password", key="prot_user")
            with col2:
                owner_pw = st.text_input("Owner password (optional)", type="password", key="prot_owner")
            if st.button("Protect PDF", use_container_width=True) and user_pw:
                with st.spinner("Protecting…"):
                    try:
                        result = protect_pdf(pdf_bytes, user_pw, owner_pw or None)
                        st.success("✅ PDF is now password protected!")
                        out_name = os.path.splitext(f.name)[0] + "_protected.pdf"
                        st.download_button("⬇ Download Protected PDF", result,
                                            out_name, "application/pdf", key="prot_dl")
                    except Exception as e:
                        st.error(f"❌ {e}")
        else:
            if not info["encrypted"]:
                st.warning("This PDF is not password-protected — there's nothing to unlock.")
            else:
                pw = st.text_input("Enter the current password", type="password", key="unlock_pw")
                if st.button("Remove Password", use_container_width=True) and pw:
                    with st.spinner("Unlocking…"):
                        try:
                            result = unlock_pdf(pdf_bytes, pw)
                            st.success("✅ Password removed! The PDF is now unlocked.")
                            out_name = os.path.splitext(f.name)[0] + "_unlocked.pdf"
                            st.download_button("⬇ Download Unlocked PDF", result,
                                                out_name, "application/pdf", key="unlock_dl")
                        except Exception as e:
                            st.error(f"❌ Incorrect password. Please check and try again.")

# ── Extract Images ────────────────────────────────────────────────
elif selected_tool == "Extract Images":
    f = st.file_uploader("Upload PDF", type=["pdf"])
    if f:
        pdf_bytes = f.read()
        info = get_pdf_info(pdf_bytes)
        col1, col2, col3 = st.columns(3)
        with col1: dpi = st.select_slider("DPI quality", [72,100,150,200,300], value=150)
        with col2: fmt = st.selectbox("Format", ["PNG","JPEG"])
        with col3: st.write(f"**{info['pages']} pages** to extract")

        if st.button("Extract Pages as Images"):
            with st.spinner(f"Rendering {info['pages']} pages…"):
                try:
                    results = extract_as_images(pdf_bytes, dpi, fmt)
                    st.success(f"✅ Extracted {len(results)} images")
                    zb = io.BytesIO()
                    with zipfile.ZipFile(zb,'w') as zf:
                        for name, data in results: zf.writestr(name, data)
                    st.download_button("⬇ Download All as ZIP", zb.getvalue(), "pages.zip", "application/zip")

                    # Preview first 3
                    st.markdown("**Preview (first 3 pages):**")
                    cols = st.columns(min(3, len(results)))
                    for i, (name, data) in enumerate(results[:3]):
                        with cols[i]:
                            st.image(data, caption=name, use_column_width=True)
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Add Text Annotation ───────────────────────────────────────────
elif selected_tool == "Add Text":

    FONTS = {
        "Serif (Arabic ✓)":      "/usr/share/fonts/truetype/freefont/FreeSerif.ttf",
        "Serif Bold (Arabic ✓)": "/usr/share/fonts/truetype/freefont/FreeSerifBold.ttf",
        "Serif Italic":          "/usr/share/fonts/truetype/freefont/FreeSerifItalic.ttf",
        "Sans (Arabic ✓)":       "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "Sans Bold (Arabic ✓)":  "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
        "Sans Italic":           "/usr/share/fonts/truetype/freefont/FreeSansOblique.ttf",
        "Mono":                  "/usr/share/fonts/truetype/freefont/FreeMono.ttf",
        "DejaVu Sans":           "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "DejaVu Sans Bold":      "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    }

    # Session state
    if "at_pdf"   not in st.session_state: st.session_state.at_pdf   = None
    if "at_queue" not in st.session_state: st.session_state.at_queue = []

    f = st.file_uploader("Upload PDF", type=["pdf"], key="at_upload")
    if f:
        data = f.read()
        if st.session_state.at_pdf != data:
            st.session_state.at_pdf   = data
            st.session_state.at_queue = []

    if st.session_state.at_pdf:
        pdf_bytes = st.session_state.at_pdf
        info = get_pdf_info(pdf_bytes)

        # ── Text styling panel ──────────────────────────────────
        st.markdown("#### 1 · Compose your text")
        text = st.text_input("Text content", "Type your text here",
                              key="at_text", label_visibility="collapsed")

        s1, s2, s3, s4, s5 = st.columns([2, 1, 1, 1, 1])
        with s1: font_name = st.selectbox("Font", list(FONTS.keys()), key="at_font")
        with s2: size  = st.number_input("Size", 6, 120, 24, key="at_size")
        with s3: color = st.color_picker("Colour", "#1c1917", key="at_color")
        with s4: page  = st.number_input("Page", 1, info["pages"], 1, key="at_page")
        with s5: align = st.selectbox("Align", ["Left", "Center", "Right"], key="at_align")

        st.markdown("#### 2 · Click on the page to place your text")

        # Render the selected page as background for the canvas
        try:
            from streamlit_drawable_canvas import st_canvas
            from PIL import ImageFont, ImageDraw

            import subprocess, tempfile, glob
            with tempfile.TemporaryDirectory() as tmp:
                pp = os.path.join(tmp, "p.pdf")
                with open(pp, "wb") as fo: fo.write(pdf_bytes)
                pref = os.path.join(tmp, "pg")
                subprocess.run(["pdftoppm", "-r", "100", "-png",
                                "-f", str(int(page)), "-l", str(int(page)),
                                pp, pref], capture_output=True, timeout=40)
                matches = glob.glob(pref + "*")
                bg_img = Image.open(matches[0]).convert("RGB") if matches else None

            if bg_img:
                # Scale canvas to fit nicely (max 700px wide)
                orig_w, orig_h = bg_img.size
                disp_w = min(700, orig_w)
                disp_h = int(disp_w * orig_h / orig_w)
                bg_resized = bg_img.resize((disp_w, disp_h))

                # Text preview image to stamp on canvas
                canvas_result = st_canvas(
                    fill_color="rgba(139, 94, 82, 0.15)",
                    stroke_width=2,
                    stroke_color="#8b5e52",
                    background_image=bg_resized,
                    update_streamlit=True,
                    height=disp_h,
                    width=disp_w,
                    drawing_mode="point",
                    point_display_radius=6,
                    key="at_canvas",
                )

                # Capture click position
                clicked_x, clicked_y = None, None
                if canvas_result.json_data and canvas_result.json_data.get("objects"):
                    last = canvas_result.json_data["objects"][-1]
                    clicked_x = last.get("left", 0)
                    clicked_y = last.get("top", 0)
                    st.caption(f"📍 Position set at ({int(clicked_x)}, {int(clicked_y)}) — click **Add** below.")

                # Store the mapping info for apply step
                st.session_state.at_disp_w = disp_w
                st.session_state.at_disp_h = disp_h

                col_add, _ = st.columns([1, 2])
                with col_add:
                    if st.button("＋ Add Text", use_container_width=True, key="at_add",
                                 disabled=(clicked_x is None)):
                        st.session_state.at_queue.append({
                            "text": text, "font": FONTS[font_name], "fname": font_name,
                            "size": int(size), "color": color, "page": int(page),
                            "align": align,
                            "cx": clicked_x, "cy": clicked_y,
                            "disp_w": disp_w, "disp_h": disp_h,
                        })
                        st.success("Added to document")
            else:
                st.error("Could not render page preview.")
        except ImportError:
            st.error("The drawable canvas component is not installed. "
                     "Add `streamlit-drawable-canvas` to requirements.txt and reboot.")

        # ── Queue & apply ───────────────────────────────────────
        if st.session_state.at_queue:
            st.markdown("#### 3 · Review & apply")
            for i, ann in enumerate(st.session_state.at_queue):
                ca, cb = st.columns([6, 1])
                with ca:
                    st.markdown(
                        f'<div class="result-card">'
                        f'<div class="fname">"{ann["text"][:55]}"</div>'
                        f'<div class="fmeta">Page {ann["page"]} · {ann["fname"]} · '
                        f'{ann["size"]}pt · {ann["align"]}</div></div>',
                        unsafe_allow_html=True)
                with cb:
                    if st.button("✕", key=f"at_del_{i}"):
                        st.session_state.at_queue.pop(i)
                        st.rerun()

            b1, b2 = st.columns(2)
            with b1:
                if st.button("🗑 Clear all", use_container_width=True, key="at_clear"):
                    st.session_state.at_queue = []
                    st.rerun()
            with b2:
                if st.button("✓ Apply & Download", use_container_width=True, key="at_apply"):
                    with st.spinner("Embedding text…"):
                        try:
                            from PIL import ImageFont, ImageDraw
                            from reportlab.lib.utils import ImageReader
                            result = pdf_bytes
                            for ann in st.session_state.at_queue:
                                # Render text → transparent PNG (Arabic-safe)
                                fp = ImageFont.truetype(ann["font"], ann["size"] * 3)
                                dummy = Image.new("RGBA", (1, 1))
                                bb = ImageDraw.Draw(dummy).textbbox((0, 0), ann["text"], font=fp)
                                tw = max(1, bb[2]-bb[0]+20)
                                th = max(1, bb[3]-bb[1]+20)
                                ti = Image.new("RGBA", (tw, th), (255, 255, 255, 0))
                                ch = ann["color"].lstrip("#")
                                cr, cg, cb2 = int(ch[0:2],16), int(ch[2:4],16), int(ch[4:6],16)
                                ImageDraw.Draw(ti).text((10, 10), ann["text"], font=fp,
                                                         fill=(cr, cg, cb2, 255))

                                # Map canvas click → PDF coords
                                r_t = PdfReader(io.BytesIO(result))
                                pg_t = r_t.pages[ann["page"]-1]
                                pdf_w = float(pg_t.mediabox.width)
                                pdf_h = float(pg_t.mediabox.height)
                                scale = pdf_w / ann["disp_w"]
                                disp_text_w = (tw/3)  # points
                                # Alignment offset
                                if ann["align"] == "Center":
                                    x_off = -disp_text_w/2
                                elif ann["align"] == "Right":
                                    x_off = -disp_text_w
                                else:
                                    x_off = 0
                                pdf_x = ann["cx"] * scale + x_off
                                pdf_y = pdf_h - (ann["cy"] * scale) - (th/6)

                                ib = io.BytesIO(); ti.save(ib, "PNG"); ib.seek(0)
                                ob = io.BytesIO()
                                oc = rl_canvas.Canvas(ob, pagesize=(pdf_w, pdf_h))
                                oc.drawImage(ImageReader(ib), pdf_x, pdf_y,
                                             width=tw/3, height=th/3, mask="auto")
                                oc.save()
                                ov = PdfReader(io.BytesIO(ob.getvalue())).pages[0]
                                rd = PdfReader(io.BytesIO(result))
                                wt = PdfWriter()
                                for ii, pg in enumerate(rd.pages):
                                    if ii == ann["page"]-1:
                                        pg.merge_page(ov)
                                    wt.add_page(pg)
                                ob2 = io.BytesIO(); wt.write(ob2)
                                result = ob2.getvalue()

                            name = f.name if f else "annotated.pdf"
                            out_name = os.path.splitext(name)[0] + "_annotated.pdf"
                            st.success(f"✅ {len(st.session_state.at_queue)} text element(s) added!")
                            st.download_button("⬇ Download PDF", result,
                                                out_name, "application/pdf", key="at_dl")
                            st.session_state.at_queue = []
                        except Exception as e:
                            st.error(f"❌ {e}")

# ── Redact ────────────────────────────────────────────────────────
elif selected_tool == "Redact":
    f = st.file_uploader("Upload PDF", type=["pdf"])
    if f:
        pdf_bytes = f.read()
        info = get_pdf_info(pdf_bytes)
        st.info("🔒 True redaction — content under the box is permanently destroyed, not just hidden. Text cannot be recovered.")

        page_num_r = st.number_input("Page number", 1, info["pages"], 1)

        col1, col2 = st.columns(2)
        with col1:
            x_pct = st.slider("← Box left edge (%)", 0, 95, 10, key="rdx")
            w_pct = st.slider("Box width (%)", 5, 100, 40, key="rdw")
        with col2:
            y_pct = st.slider("↑ Box top edge (%)", 0, 95, 10, key="rdy")
            h_pct = st.slider("Box height (%)", 2, 50, 5, key="rdh")

        # Live preview
        try:
            import subprocess, tempfile, glob
            from PIL import Image as PILImg, ImageDraw
            with tempfile.TemporaryDirectory() as tmp:
                pp = os.path.join(tmp, "p.pdf")
                with open(pp, "wb") as fo: fo.write(pdf_bytes)
                pref = os.path.join(tmp, "pg")
                subprocess.run(["pdftoppm","-r","90","-png","-f",str(page_num_r),
                                "-l",str(page_num_r),pp,pref], capture_output=True, timeout=30)
                m = glob.glob(pref+"*")
                if m:
                    prev = PILImg.open(m[0]).convert("RGB")
                    pw2, ph2 = prev.size
                    draw = ImageDraw.Draw(prev)
                    bl = x_pct/100*pw2
                    bt = y_pct/100*ph2
                    br = min(pw2, bl + w_pct/100*pw2)
                    bb = min(ph2, bt + h_pct/100*ph2)
                    draw.rectangle([bl,bt,br,bb], fill=(0,0,0))
                    st.image(prev, caption=f"Preview — Page {page_num_r} (black = redacted)", use_column_width=True)
        except Exception:
            pass

        if st.button("Apply Redaction", use_container_width=True):
            with st.spinner("Applying permanent redaction…"):
                try:
                    # Convert % to PDF points
                    reader = PdfReader(io.BytesIO(pdf_bytes))
                    pw = float(reader.pages[page_num_r-1].mediabox.width)
                    ph = float(reader.pages[page_num_r-1].mediabox.height)
                    x_pt = x_pct/100 * pw
                    w_pt = w_pct/100 * pw
                    h_pt = h_pct/100 * ph
                    # y from bottom = ph - top - height
                    y_pt = ph - (y_pct/100 * ph) - h_pt
                    result = redact_pdf(pdf_bytes, page_num_r, x_pt, y_pt, w_pt, h_pt)
                    st.success(f"✅ Page {page_num_r} redacted — content permanently removed")
                    out_name = os.path.splitext(f.name)[0] + "_redacted.pdf"
                    st.download_button("⬇ Download", result, out_name, "application/pdf")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Reorder Pages ─────────────────────────────────────────────────
elif selected_tool == "Reorder Pages":
    f = st.file_uploader("Upload PDF", type=["pdf"])
    if f:
        pdf_bytes = f.read()
        info = get_pdf_info(pdf_bytes)
        st.info(f"📄 {info['pages']} pages — enter the new page order below.")
        default_order = ", ".join(str(i) for i in range(1, info['pages']+1))
        order_input = st.text_input("New page order (comma separated)", value=default_order)
        st.caption("Example: '3, 1, 2' moves page 3 to the front.")

        if st.button("Reorder Pages"):
            with st.spinner("Reordering…"):
                try:
                    new_order = [int(x.strip()) for x in order_input.split(',') if x.strip().isdigit()]
                    result = reorder_pages(pdf_bytes, new_order)
                    st.success(f"✅ Reordered to: {new_order}")
                    out_name = os.path.splitext(f.name)[0] + "_reordered.pdf"
                    st.download_button("⬇ Download", result, out_name, "application/pdf")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Images → PDF ──────────────────────────────────────────────────
elif selected_tool == "Images → PDF":
    st.markdown("Convert JPG, PNG, WEBP, BMP or TIFF images into a single PDF.")
    files = st.file_uploader(
        "Upload images (they will appear in the PDF in the order uploaded)",
        type=["jpg","jpeg","png","webp","bmp","tiff","tif"],
        accept_multiple_files=True, key="img2pdf_files"
    )
    if files:
        col1, col2 = st.columns(2)
        with col1:
            page_size = st.selectbox("Page size", ["A4","Letter","Fit to image"], key="img2pdf_size")
            orientation = st.selectbox("Orientation", ["Portrait","Landscape"], key="img2pdf_orient")
        with col2:
            margin_mm = st.slider("Margin (mm)", 0, 40, 10, key="img2pdf_margin")
            out_name_img = st.text_input("Output filename", value="images.pdf", key="img2pdf_name")

        # Show thumbnails
        st.markdown(f"**{len(files)} image(s) selected:**")
        thumb_cols = st.columns(min(6, len(files)))
        for i, imgf in enumerate(files[:6]):
            with thumb_cols[i]:
                st.image(imgf, use_column_width=True)
                st.caption(imgf.name[:16])

        if st.button("Convert to PDF", key="img2pdf_btn", use_container_width=True):
            with st.spinner(f"Converting {len(files)} image(s)…"):
                try:
                    from PIL import Image as PILImage
                    from reportlab.lib.pagesizes import A4, letter
                    from reportlab.lib.utils import ImageReader

                    PAGE_SIZES = {"A4": A4, "Letter": letter}
                    margin_pt = margin_mm * 2.835  # mm to points

                    pil_images = []
                    for imgf in files:
                        img = PILImage.open(imgf).convert("RGB")
                        pil_images.append(img)

                    if page_size == "Fit to image":
                        # Each page fits its image
                        first_w, first_h = pil_images[0].size
                        ps = (first_w, first_h)
                    else:
                        ps = PAGE_SIZES.get(page_size, A4)
                        if orientation == "Landscape":
                            ps = (ps[1], ps[0])

                    out_buf = io.BytesIO()
                    c = rl_canvas.Canvas(out_buf, pagesize=ps)

                    for idx, img in enumerate(pil_images):
                        if page_size == "Fit to image":
                            iw, ih = img.size
                            c.setPageSize((iw, ih))
                            avail_w, avail_h = iw - margin_pt*2, ih - margin_pt*2
                        else:
                            avail_w = ps[0] - margin_pt*2
                            avail_h = ps[1] - margin_pt*2

                        iw, ih = img.size
                        scale = min(avail_w/iw, avail_h/ih)
                        draw_w, draw_h = iw*scale, ih*scale
                        x = margin_pt + (avail_w - draw_w)/2
                        y = margin_pt + (avail_h - draw_h)/2

                        img_buf = io.BytesIO()
                        img.save(img_buf, "PNG")
                        img_buf.seek(0)
                        c.drawImage(ImageReader(img_buf), x, y, width=draw_w, height=draw_h)
                        c.showPage()

                    c.save()
                    pdf_data = out_buf.getvalue()
                    st.success(f"✅ {len(pil_images)} images → {len(pil_images)}-page PDF")
                    st.download_button("⬇ Download PDF", pdf_data,
                                        out_name_img, "application/pdf", key="img2pdf_dl")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Word → PDF ────────────────────────────────────────────────────
elif selected_tool == "Word → PDF":
    st.markdown("Convert `.docx` Word documents to PDF using LibreOffice.")
    files = st.file_uploader("Upload Word documents (.docx)",
                              type=["docx"], accept_multiple_files=True, key="word2pdf_files")
    if files:
        as_zip_w = st.checkbox("Download all as ZIP", value=len(files)>1, key="word2pdf_zip")
        if st.button("Convert to PDF", use_container_width=True, key="word2pdf_btn"):
            import tempfile, subprocess
            results = []
            bar = st.progress(0)
            for idx, f in enumerate(files):
                with st.spinner(f"Converting {f.name}…"):
                    try:
                        with tempfile.TemporaryDirectory() as tmp:
                            in_path = os.path.join(tmp, f.name)
                            with open(in_path, "wb") as fout:
                                fout.write(f.read())
                            r = subprocess.run(
                                ["libreoffice","--headless","--convert-to","pdf",
                                 "--outdir", tmp, in_path],
                                capture_output=True, text=True, timeout=60
                            )
                            pdf_path = in_path.replace(".docx",".pdf")
                            if os.path.exists(pdf_path):
                                with open(pdf_path,"rb") as pf:
                                    pdf_data = pf.read()
                                out_n = os.path.splitext(f.name)[0]+".pdf"
                                results.append({"name":f.name,"out":out_n,"bytes":pdf_data,"ok":True})
                            else:
                                results.append({"name":f.name,"error":r.stderr[:200],"ok":False})
                    except Exception as e:
                        results.append({"name":f.name,"error":str(e),"ok":False})
                bar.progress((idx+1)/len(files))
            bar.empty()

            if as_zip_w and any(r["ok"] for r in results):
                zb = io.BytesIO()
                with zipfile.ZipFile(zb,"w",zipfile.ZIP_DEFLATED) as zf:
                    for r in results:
                        if r["ok"]: zf.writestr(r["out"], r["bytes"])
                st.download_button("⬇ Download ZIP", zb.getvalue(),
                                    "converted.zip", "application/zip", key="word2pdf_zip_dl")
            for r in results:
                status = "✓ Converted" if r["ok"] else f"✗ {r.get('error','')}"
                color = "#6b8b5e" if r["ok"] else "#8b5e5e"
                st.markdown(
                    f'<div class="result-card"><div class="fname">📄 {r["name"]}</div>'
                    f'<div class="fmeta" style="color:{color}">{status}</div></div>',
                    unsafe_allow_html=True)
                if r["ok"]:
                    st.download_button(f"⬇ Download {r['out']}", r["bytes"], r["out"],
                                        "application/pdf", key=f"w2p_{r['out']}")

# ── Excel → PDF ───────────────────────────────────────────────────
elif selected_tool == "Excel → PDF":
    st.markdown("Convert `.xlsx` Excel spreadsheets to PDF using LibreOffice.")
    files = st.file_uploader("Upload Excel files (.xlsx)",
                              type=["xlsx","xls"], accept_multiple_files=True, key="xl2pdf_files")
    if files:
        as_zip_x = st.checkbox("Download all as ZIP", value=len(files)>1, key="xl2pdf_zip")
        if st.button("Convert to PDF", use_container_width=True, key="xl2pdf_btn"):
            import tempfile, subprocess
            results = []
            bar2 = st.progress(0)
            for idx, f in enumerate(files):
                with st.spinner(f"Converting {f.name}…"):
                    try:
                        with tempfile.TemporaryDirectory() as tmp:
                            in_path = os.path.join(tmp, f.name)
                            with open(in_path,"wb") as fout:
                                fout.write(f.read())
                            r = subprocess.run(
                                ["libreoffice","--headless","--convert-to","pdf",
                                 "--outdir", tmp, in_path],
                                capture_output=True, text=True, timeout=60
                            )
                            ext = os.path.splitext(f.name)[1]
                            pdf_path = in_path.replace(ext, ".pdf")
                            if os.path.exists(pdf_path):
                                with open(pdf_path,"rb") as pf:
                                    pdf_data = pf.read()
                                out_n = os.path.splitext(f.name)[0]+".pdf"
                                results.append({"name":f.name,"out":out_n,"bytes":pdf_data,"ok":True})
                            else:
                                results.append({"name":f.name,"error":r.stderr[:200],"ok":False})
                    except Exception as e:
                        results.append({"name":f.name,"error":str(e),"ok":False})
                bar2.progress((idx+1)/len(files))
            bar2.empty()

            if as_zip_x and any(r["ok"] for r in results):
                zb2 = io.BytesIO()
                with zipfile.ZipFile(zb2,"w",zipfile.ZIP_DEFLATED) as zf2:
                    for r in results:
                        if r["ok"]: zf2.writestr(r["out"], r["bytes"])
                st.download_button("⬇ Download ZIP", zb2.getvalue(),
                                    "converted_xl.zip", "application/zip", key="xl2pdf_zip_dl")
            for r in results:
                status = "✓ Converted" if r["ok"] else f"✗ {r.get('error','')}"
                color = "#6b8b5e" if r["ok"] else "#8b5e5e"
                st.markdown(
                    f'<div class="result-card"><div class="fname">📊 {r["name"]}</div>'
                    f'<div class="fmeta" style="color:{color}">{status}</div></div>',
                    unsafe_allow_html=True)
                if r["ok"]:
                    st.download_button(f"⬇ Download {r['out']}", r["bytes"], r["out"],
                                        "application/pdf", key=f"x2p_{r['out']}")

# ── PDF → Excel ───────────────────────────────────────────────────
elif selected_tool == "PDF → Excel":
    f = st.file_uploader("Upload PDF", type=["pdf"], key="pdf2xl")
    if f:
        pdf_bytes = f.read()
        st.info("Tables detected in the PDF will be extracted as Excel sheets. Text paragraphs go to a 'Text' sheet.")
        if st.button("Convert to Excel", use_container_width=True, key="pdf2xl_btn"):
            with st.spinner("Extracting…"):
                try:
                    import openpyxl
                    from openpyxl.styles import Font, PatternFill, Alignment
                    wb = openpyxl.Workbook()
                    text_ws = wb.active
                    text_ws.title = "Text"
                    text_row = 1

                    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                        for p_num, page in enumerate(pdf.pages, 1):
                            # Tables → individual sheets
                            tables = page.extract_tables()
                            for t_idx, table in enumerate(tables):
                                ws = wb.create_sheet(f"Page{p_num}_Table{t_idx+1}")
                                for r_idx, row in enumerate(table):
                                    for c_idx, cell in enumerate(row):
                                        c = ws.cell(r_idx+1, c_idx+1, cell or "")
                                        if r_idx == 0:
                                            c.font = Font(bold=True)
                                            c.fill = PatternFill("solid", fgColor="E8E0D5")

                            # Text → Text sheet
                            raw = page.extract_text()
                            if raw:
                                text_ws.cell(text_row, 1, f"— Page {p_num} —").font = Font(bold=True, color="8B5E52")
                                text_row += 1
                                for line in raw.splitlines():
                                    if line.strip():
                                        text_ws.cell(text_row, 1, line.strip())
                                        text_row += 1
                                text_row += 1

                    text_ws.column_dimensions["A"].width = 80
                    out = io.BytesIO(); wb.save(out)
                    out_name = os.path.splitext(f.name)[0]+".xlsx"
                    st.success("✅ Converted!")
                    st.download_button("⬇ Download Excel", out.getvalue(),
                                        out_name, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                        key="pdf2xl_dl")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── PDF → CSV ─────────────────────────────────────────────────────
elif selected_tool == "PDF → CSV":
    f = st.file_uploader("Upload PDF", type=["pdf"], key="pdf2csv")
    if f:
        pdf_bytes = f.read()
        st.info("All tables from the PDF will be combined into one CSV file.")
        if st.button("Convert to CSV", use_container_width=True, key="pdf2csv_btn"):
            with st.spinner("Extracting tables…"):
                try:
                    import csv as csv_mod
                    rows_out = []
                    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                        for p_num, page in enumerate(pdf.pages, 1):
                            tables = page.extract_tables()
                            for table in tables:
                                if rows_out:
                                    rows_out.append([])  # blank row between tables
                                rows_out.extend([[c or "" for c in row] for row in table])

                    if not rows_out:
                        st.warning("No tables found. Falling back to plain text extraction.")
                        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                            for page in pdf.pages:
                                raw = page.extract_text()
                                if raw:
                                    for line in raw.splitlines():
                                        rows_out.append([line])

                    out = io.StringIO()
                    csv_mod.writer(out).writerows(rows_out)
                    out_name = os.path.splitext(f.name)[0]+".csv"
                    st.success(f"✅ {len(rows_out)} rows extracted!")
                    st.download_button("⬇ Download CSV", out.getvalue().encode(),
                                        out_name, "text/csv", key="pdf2csv_dl")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── PDF → Images ─────────────────────────────────────────────────
elif selected_tool == "PDF → Images":
    f = st.file_uploader("Upload PDF", type=["pdf"], key="pdf2img")
    if f:
        pdf_bytes = f.read()
        info = get_pdf_info(pdf_bytes)

        col1, col2, col3 = st.columns(3)
        with col1:
            dpi = st.select_slider("Quality (DPI)", [72, 100, 120], 72, key="pdf2img_dpi")
        with col2:
            fmt = st.selectbox("Format", ["JPEG", "PNG"], key="pdf2img_fmt")
        with col3:
            st.metric("Pages", info["pages"])

        page_range = st.text_input(
            "Page range (e.g. 1-5, or blank for all)",
            value="1-5" if info["pages"] > 5 else "",
            key="pdf2img_range"
        )
        st.caption("Free tier tip: convert max 5 pages at a time to avoid memory limits.")

        if st.button("Convert to Images", use_container_width=True, key="pdf2img_btn"):
            import subprocess, tempfile, glob

            try:
                if page_range.strip():
                    parts = page_range.strip().split("-")
                    p_start = max(1, int(parts[0]))
                    p_end   = min(info["pages"], int(parts[1]) if len(parts) > 1 else p_start)
                else:
                    p_start, p_end = 1, info["pages"]
            except Exception:
                st.error("Invalid page range. Use format like 1-5.")
                p_start, p_end = 1, 1

            total = p_end - p_start + 1
            bar    = st.progress(0)
            status = st.empty()
            previews = []
            zb = io.BytesIO()

            try:
                with tempfile.TemporaryDirectory() as tmp:
                    pdf_path = os.path.join(tmp, "input.pdf")
                    with open(pdf_path, "wb") as fout:
                        fout.write(pdf_bytes)

                    with zipfile.ZipFile(zb, "w", zipfile.ZIP_DEFLATED) as zf:
                        for page_num in range(p_start, p_end + 1):
                            status.text(f"Rendering page {page_num}/{p_end}…")
                            out_prefix = os.path.join(tmp, f"p{page_num:03d}")
                            flag = "-jpeg" if fmt == "JPEG" else "-png"
                            subprocess.run(
                                ["pdftoppm", "-r", str(dpi), flag,
                                 "-f", str(page_num), "-l", str(page_num),
                                 pdf_path, out_prefix],
                                capture_output=True, timeout=60
                            )
                            # Find output file
                            matches = glob.glob(out_prefix + "*")
                            if matches:
                                img_path = matches[0]
                                ext = "jpg" if fmt == "JPEG" else "png"
                                with open(img_path, "rb") as img_f:
                                    img_bytes = img_f.read()
                                zf.writestr(f"page_{page_num:03d}.{ext}", img_bytes)
                                if len(previews) < 3:
                                    from PIL import Image as PILImg
                                    previews.append((page_num,
                                        PILImg.open(io.BytesIO(img_bytes)).copy()))
                                os.remove(img_path)
                            bar.progress((page_num - p_start + 1) / total)

                bar.empty()
                status.empty()
                st.success(f"✅ {total} page(s) converted!")
                st.download_button(
                    "⬇ Download ZIP", zb.getvalue(),
                    os.path.splitext(f.name)[0] + "_images.zip",
                    "application/zip", key="pdf2img_dl"
                )
                if previews:
                    st.markdown("**Preview:**")
                    cols = st.columns(len(previews))
                    for col, (pnum, img) in zip(cols, previews):
                        with col:
                            st.image(img, caption=f"Page {pnum}", use_column_width=True)

            except Exception as e:
                bar.empty()
                status.empty()
                st.error(f"❌ {e}")

# ── PDF → Text ────────────────────────────────────────────────────
elif selected_tool == "PDF → Text":
    f = st.file_uploader("Upload PDF", type=["pdf"], key="pdf2txt")
    if f:
        pdf_bytes = f.read()
        info = get_pdf_info(pdf_bytes)
        col1, col2 = st.columns(2)
        with col1: ocr_on = st.checkbox("OCR for scanned PDFs", True, key="pdf2txt_ocr")
        with col2: page_breaks = st.checkbox("Add page separators", True, key="pdf2txt_pb")

        if st.button("Extract Text", use_container_width=True, key="pdf2txt_btn"):
            with st.spinner("Extracting text…"):
                try:
                    all_text = []
                    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                        for p_num, page in enumerate(pdf.pages, 1):
                            raw = page.extract_text() or ""
                            if not raw.strip() and ocr_on:
                                try:
                                    import pytesseract
                                    imgs_ocr = convert_from_bytes(pdf_bytes, dpi=150,
                                                                   first_page=p_num, last_page=p_num)
                                    raw = pytesseract.image_to_string(imgs_ocr[0])
                                except Exception:
                                    pass
                            if page_breaks:
                                sep = "─"*60
                                all_text.append(f"\n{sep}\nPage {p_num}\n{sep}\n")
                            all_text.append(raw)

                    full_text = "\n".join(all_text)
                    st.success(f"✅ {len(full_text):,} characters extracted!")
                    st.text_area("Extracted text (preview)", full_text[:3000]+
                                  ("…" if len(full_text)>3000 else ""), height=300)
                    out_name = os.path.splitext(f.name)[0]+".txt"
                    st.download_button("⬇ Download .txt", full_text.encode("utf-8"),
                                        out_name, "text/plain", key="pdf2txt_dl")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── PDF → Markdown ────────────────────────────────────────────────
elif selected_tool == "PDF → Markdown":
    f = st.file_uploader("Upload PDF", type=["pdf"], key="pdf2md")
    if f:
        pdf_bytes = f.read()
        info = get_pdf_info(pdf_bytes)
        include_tables = st.checkbox("Convert tables to Markdown format", True, key="pdf2md_tables")

        if st.button("Convert to Markdown", use_container_width=True, key="pdf2md_btn"):
            with st.spinner("Converting…"):
                try:
                    md_lines = []
                    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                        for p_num, page in enumerate(pdf.pages, 1):
                            if p_num > 1:
                                md_lines.append(f"\n---\n")

                            # Tables first
                            if include_tables:
                                tables = page.extract_tables()
                                table_bboxes = []
                                try:
                                    table_bboxes = [t.bbox for t in page.find_tables()]
                                except Exception:
                                    pass
                                for table in tables:
                                    if not table: continue
                                    rows = [[str(c or "").strip() for c in row] for row in table]
                                    if not rows: continue
                                    ncols = max(len(r) for r in rows)
                                    # Pad rows
                                    rows = [r + [""]*(ncols-len(r)) for r in rows]
                                    md_lines.append("| " + " | ".join(rows[0]) + " |")
                                    md_lines.append("| " + " | ".join(["---"]*ncols) + " |")
                                    for row in rows[1:]:
                                        md_lines.append("| " + " | ".join(row) + " |")
                                    md_lines.append("")

                            # Text with heading detection
                            words = page.extract_words(extra_attrs=["size", "fontname"])
                            if not words:
                                raw = page.extract_text() or ""
                                for line in raw.splitlines():
                                    if line.strip():
                                        md_lines.append(line.strip())
                                continue

                            lines = {}
                            for w in words:
                                lines.setdefault(round(w["top"], 1), []).append(w)

                            def in_table_bbox(y):
                                for bbox in table_bboxes:
                                    if bbox[1] <= y <= bbox[3]:
                                        return True
                                return False

                            prev_bottom = None
                            for top, lw in sorted(lines.items()):
                                if in_table_bbox(top):
                                    continue
                                lw.sort(key=lambda w: w["x0"])
                                text = " ".join(w["text"] for w in lw).strip()
                                if not text:
                                    continue
                                sizes = [w.get("size", 0) for w in lw if w.get("size")]
                                fs = max(sizes) if sizes else 12
                                bold = any("Bold" in w.get("fontname", "") for w in lw)
                                # Blank line for paragraph gap
                                if prev_bottom and (top - prev_bottom) > 14:
                                    md_lines.append("")
                                if fs >= 18:
                                    md_lines.append(f"# {text}")
                                elif fs >= 14:
                                    md_lines.append(f"## {text}")
                                elif fs >= 12 and bold:
                                    md_lines.append(f"### {text}")
                                elif bold:
                                    md_lines.append(f"**{text}**")
                                else:
                                    md_lines.append(text)
                                prev_bottom = top + lw[0].get("height", 10)

                    md_out = "\n".join(md_lines)
                    out_name = os.path.splitext(f.name)[0] + ".md"
                    st.success(f"✅ Converted — {len(md_out):,} characters")
                    st.text_area("Preview", md_out[:2000] + ("…" if len(md_out) > 2000 else ""), height=300)
                    st.download_button("⬇ Download .md", md_out.encode("utf-8"),
                                        out_name, "text/markdown", key="pdf2md_dl")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Word → Markdown ───────────────────────────────────────────────
elif selected_tool == "Word → Markdown":
    f = st.file_uploader("Upload Word document (.docx)", type=["docx"], key="docx2md")
    if f:
        if st.button("Convert to Markdown", use_container_width=True, key="docx2md_btn"):
            with st.spinner("Converting…"):
                try:
                    from docx import Document as DocxDoc
                    from docx.oxml.ns import qn

                    doc = DocxDoc(io.BytesIO(f.read()))
                    md_lines = []

                    for para in doc.paragraphs:
                        style = para.style.name
                        # Build text preserving inline bold/italic
                        inline = ""
                        for run in para.runs:
                            txt = run.text
                            if not txt:
                                continue
                            if run.bold and run.italic:
                                txt = f"***{txt}***"
                            elif run.bold:
                                txt = f"**{txt}**"
                            elif run.italic:
                                txt = f"*{txt}*"
                            inline += txt

                        inline = inline.strip()
                        if not inline:
                            md_lines.append("")
                            continue

                        if "Title" in style:
                            md_lines.append(f"# {inline}\n")
                        elif "Heading 1" in style:
                            md_lines.append(f"# {inline}\n")
                        elif "Heading 2" in style:
                            md_lines.append(f"## {inline}\n")
                        elif "Heading 3" in style:
                            md_lines.append(f"### {inline}\n")
                        elif "Heading 4" in style:
                            md_lines.append(f"#### {inline}\n")
                        elif "List" in style or para.style.name.startswith("List"):
                            md_lines.append(f"- {inline}")
                        else:
                            md_lines.append(f"{inline}\n")

                    # Tables
                    for table in doc.tables:
                        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                        if not rows: continue
                        ncols = max(len(r) for r in rows)
                        rows = [r + [""]*(ncols-len(r)) for r in rows]
                        md_lines.append("")
                        md_lines.append("| " + " | ".join(rows[0]) + " |")
                        md_lines.append("| " + " | ".join(["---"]*ncols) + " |")
                        for row in rows[1:]:
                            md_lines.append("| " + " | ".join(row) + " |")
                        md_lines.append("")

                    md_out = "\n".join(md_lines)
                    out_name = os.path.splitext(f.name)[0] + ".md"
                    st.success(f"✅ Converted — {len(md_out):,} characters")
                    st.text_area("Preview", md_out[:2000] + ("…" if len(md_out) > 2000 else ""), height=300)
                    st.download_button("⬇ Download .md", md_out.encode("utf-8"),
                                        out_name, "text/markdown", key="docx2md_dl")
                except Exception as e:
                    st.error(f"❌ {e}")

# ── Excel → Markdown ──────────────────────────────────────────────
elif selected_tool == "Excel → Markdown":
    f = st.file_uploader("Upload Excel file (.xlsx)", type=["xlsx", "xls"], key="xl2md")
    if f:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(f.read()), data_only=True)
        sheet_names = wb.sheetnames
        sheet = st.selectbox("Sheet", sheet_names, key="xl2md_sheet")
        include_header = st.checkbox("First row is header", True, key="xl2md_header")

        if st.button("Convert to Markdown", use_container_width=True, key="xl2md_btn"):
            with st.spinner("Converting…"):
                try:
                    ws = wb[sheet]
                    rows = []
                    for row in ws.iter_rows(values_only=True):
                        if any(v is not None for v in row):
                            rows.append([str(v) if v is not None else "" for v in row])

                    if not rows:
                        st.warning("No data found in this sheet.")
                    else:
                        md_lines = []
                        ncols = max(len(r) for r in rows)
                        rows = [r + [""]*(ncols-len(r)) for r in rows]

                        if include_header:
                            md_lines.append("| " + " | ".join(rows[0]) + " |")
                            md_lines.append("| " + " | ".join(["---"]*ncols) + " |")
                            data_rows = rows[1:]
                        else:
                            # Auto-generate column headers
                            cols = [f"Col {i+1}" for i in range(ncols)]
                            md_lines.append("| " + " | ".join(cols) + " |")
                            md_lines.append("| " + " | ".join(["---"]*ncols) + " |")
                            data_rows = rows

                        for row in data_rows:
                            md_lines.append("| " + " | ".join(row) + " |")

                        md_out = "\n".join(md_lines)
                        out_name = os.path.splitext(f.name)[0] + f"_{sheet}.md"
                        st.success(f"✅ {len(data_rows)} rows converted from sheet '{sheet}'")
                        st.text_area("Preview", md_out[:2000] + ("…" if len(md_out) > 2000 else ""), height=300)
                        st.download_button("⬇ Download .md", md_out.encode("utf-8"),
                                            out_name, "text/markdown", key="xl2md_dl")
                except Exception as e:
                    st.error(f"❌ {e}")
