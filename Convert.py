"""
PDF to Word Converter
---------------------
Converts PDF files to clean .docx Word documents.
Preserves headings, paragraphs, and tables where possible.

Usage:
    python convert.py input.pdf                     → creates input.docx
    python convert.py input.pdf -o output.docx      → custom output name
    python convert.py *.pdf                         → batch convert multiple files
"""

import argparse
import os
import re
import sys

import pdfplumber
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def looks_like_heading(text: str, font_size: float = None, is_bold: bool = False) -> int:
    """Return heading level (1-3) or 0 if not a heading."""
    stripped = text.strip()
    if not stripped:
        return 0

    # Very short ALL CAPS lines are likely headings
    if stripped.isupper() and 3 <= len(stripped) <= 80:
        return 1

    # Bold short lines
    if is_bold and len(stripped) <= 80 and not stripped.endswith('.'):
        return 2

    # Font size heuristic
    if font_size:
        if font_size >= 18:
            return 1
        if font_size >= 14:
            return 2
        if font_size >= 12 and is_bold:
            return 3

    return 0


def clean_text(text: str) -> str:
    """Remove control characters and normalise whitespace."""
    if not text:
        return ""
    # Replace non-breaking spaces etc.
    text = text.replace('\xa0', ' ')
    # Collapse multiple spaces (but keep intentional indentation)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()


def extract_page_content(page):
    """
    Extract structured content (text blocks + tables) from a single page.
    Returns a list of dicts: {type: 'paragraph'|'table', ...}
    """
    content = []

    # Get table bounding boxes so we can skip them in regular text extraction
    tables = page.extract_tables()
    table_bboxes = [t.bbox for t in page.find_tables()] if hasattr(page, 'find_tables') else []

    # ---- TEXT ----
    words = page.extract_words(extra_attrs=["size", "fontname"])
    if not words:
        raw = page.extract_text()
        if raw:
            for line in raw.splitlines():
                line = clean_text(line)
                if line:
                    content.append({"type": "paragraph", "text": line,
                                    "font_size": None, "bold": False})
        # Add tables even if no text words
        for tbl in tables:
            if tbl:
                content.append({"type": "table", "data": tbl})
        return content

    # Group words into lines by their top-coordinate
    lines = {}
    for w in words:
        key = round(w['top'], 1)
        lines.setdefault(key, []).append(w)

    # Sort lines top-to-bottom
    sorted_lines = sorted(lines.items())

    # Track which lines overlap table bboxes
    def in_table(y):
        for bbox in table_bboxes:
            if bbox[1] <= y <= bbox[3]:
                return True
        return False

    prev_bottom = None
    for top, line_words in sorted_lines:
        if in_table(top):
            continue

        line_words.sort(key=lambda w: w['x0'])
        line_text = clean_text(' '.join(w['text'] for w in line_words))
        if not line_text:
            continue

        # Determine predominant font size and boldness for the line
        sizes = [w.get('size', 0) for w in line_words if w.get('size')]
        font_size = max(sizes) if sizes else None
        fontnames = [w.get('fontname', '') for w in line_words]
        is_bold = any('Bold' in fn or 'bold' in fn for fn in fontnames)

        # Detect blank gap between previous line (paragraph break)
        gap = (top - prev_bottom) if prev_bottom is not None else 0
        if gap > 12:
            content.append({"type": "blank"})

        content.append({
            "type": "paragraph",
            "text": line_text,
            "font_size": font_size,
            "bold": is_bold,
        })
        prev_bottom = top + (line_words[0].get('height', 10))

    # ---- TABLES ----
    for tbl in tables:
        if tbl:
            content.append({"type": "table", "data": tbl})

    return content


def add_table_to_doc(doc: Document, table_data: list):
    """Add a pdfplumber table to the Word document."""
    if not table_data:
        return
    rows = [r for r in table_data if any(c for c in r)]
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.style = 'Table Grid'

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = tbl.cell(r_idx, c_idx)
                cell.text = clean_text(cell_text or "")
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


def convert_pdf_to_word(pdf_path: str, output_path: str = None) -> str:
    """Convert a PDF file to a Word document. Returns the output path."""
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    if output_path is None:
        base = os.path.splitext(pdf_path)[0]
        output_path = base + ".docx"

    doc = Document()

    # Set reasonable margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, start=1):
            print(f"  Processing page {page_num}/{total_pages}…", end='\r')
            content = extract_page_content(page)

            for item in content:
                if item["type"] == "blank":
                    # Small spacer — skip to avoid excessive blank lines
                    continue

                elif item["type"] == "paragraph":
                    text = item["text"]
                    font_size = item.get("font_size")
                    bold = item.get("bold", False)

                    level = looks_like_heading(text, font_size, bold)

                    if level == 1:
                        doc.add_heading(text, level=1)
                    elif level == 2:
                        doc.add_heading(text, level=2)
                    elif level == 3:
                        doc.add_heading(text, level=3)
                    else:
                        p = doc.add_paragraph(text)
                        if bold:
                            for run in p.runs:
                                run.bold = True

                elif item["type"] == "table":
                    add_table_to_doc(doc, item["data"])

            # Page break between PDF pages (skip after last page)
            if page_num < total_pages:
                doc.add_page_break()

    doc.save(output_path)
    print(f"\n✅  Saved: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Convert PDF files to clean Word (.docx) documents."
    )
    parser.add_argument("pdfs", nargs="+", help="PDF file(s) to convert")
    parser.add_argument("-o", "--output", help="Output .docx path (single file only)")
    args = parser.parse_args()

    if args.output and len(args.pdfs) > 1:
        print("Error: --output can only be used when converting a single PDF.")
        sys.exit(1)

    success, failed = 0, 0
    for pdf_path in args.pdfs:
        print(f"\n📄  Converting: {pdf_path}")
        try:
            out = convert_pdf_to_word(pdf_path, args.output)
            success += 1
        except Exception as e:
            print(f"❌  Failed: {pdf_path} — {e}")
            failed += 1

    print(f"\nDone. {success} converted, {failed} failed.")


if __name__ == "__main__":
    main()
